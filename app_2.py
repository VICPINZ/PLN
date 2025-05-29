import os
import joblib
import pandas as pd
import streamlit as st
from docx import Document
import string
from collections import Counter
from sklearn.preprocessing import LabelEncoder
from sklearn.pipeline import Pipeline
import nltk
import google.generativeai as genai
from entrenar_modelo import entrenar_modelo

# Descargar stopwords de NLTK
nltk.download('stopwords')
from nltk.corpus import stopwords
stopwords_es = set(stopwords.words('spanish'))

# Configurar clave de API de Gemini
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-1.5-flash")

# --- Función para predecir clasificación ---
def predecir_clasificacion(textos):
    modelo_path = 'modelo_entrenado.pkl'
    encoder_path = 'label_encoder.pkl'

    if not os.path.isfile(modelo_path) or not os.path.isfile(encoder_path):
        st.error(f"❌ No se encontró el archivo del modelo o del codificador.\n\n"
                 f"Asegúrate de subir 'modelo_rf.pkl' y 'label_encoder.pkl' a tu repositorio y que estén en la misma carpeta que este archivo.")
        st.stop()

    pipeline = joblib.load(modelo_path)
    le = joblib.load(encoder_path)
    predicciones = pipeline.predict(textos)
    return le.inverse_transform(predicciones)

# --- Función para generar recomendaciones ---
def generar_recomendaciones(comentarios):
    entrada = " ".join(comentarios)
    prompt = f"Basado en los siguientes comentarios de recomendación, sugiere mejoras concretas:\n{entrada}\n\nRecomendaciones:"
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Ocurrió un error al generar las recomendaciones: {e}"

# --- Función para generar informe Word ---
def generar_informe(comentarios, recomendaciones, nombre_archivo, resumen):
    doc = Document()
    doc.add_heading("Informe de Resultados de la Encuesta", 0)

    doc.add_heading("Resumen por clase", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Clase'
    hdr_cells[1].text = 'Conteo'
    hdr_cells[2].text = 'Top 10 palabras'
    for fila in resumen:
        row_cells = table.add_row().cells
        row_cells[0].text = str(fila['Clase'])
        row_cells[1].text = str(fila['Conteo'])
        row_cells[2].text = fila['Top 10 palabras']

    doc.add_heading("Propuesta de Mejora", level=1)
    doc.add_paragraph(recomendaciones)
    doc.save(nombre_archivo)

# --- Interfaz Streamlit ---
st.set_page_config(page_title="Análisis de Encuesta", layout="wide")
st.title("Análisis de Encuesta - Ministerio de Defensa")

tab1, tab2, tab3 = st.tabs(["💡 Problemática Identificada","📊 Entrenamiento del Modelo", "📄 Generación de Informe"])

# --- Pestaña 1: Problema ---
with tab1:
    with st.expander("📄 Procedimiento - Análisis de Encuesta"):
        st.markdown("""
        ## PROCEDIMIENTO  
        ### APLICACIÓN Y ANÁLISIS ENCUESTA DE SATISFACCIÓN

        **Objetivo:**  
        Implementación de una encuesta con preguntas abiertas y cerradas que permita detectar el nivel de satisfacción de los usuarios y aspectos a mejorar.

        ---
        ### Proceso:

        | No | Tarea dentro del proceso o actividad | Texto(s) de entrada | Función(es) básica(s) de NLP a utilizar |
        |----|------------------------------------|--------------------|----------------------------------------|
        | 1  | Diseño de la encuesta | Temas clave de evaluación | Generación automática de texto, Clasificación de textos |
        | 2  | Implementar encuesta a participantes | Encuesta diseñada | Ninguna |
        | 3  | Recolectar datos | Respuestas de los participantes | Ninguna |
        | 4  | Separar preguntas abiertas y preguntas cerradas | Respuestas de la encuesta | Clasificación de textos |
        | 5  | Generar análisis descriptivo de preguntas cerradas | Respuestas estructuradas | Similitud de textos, Extracción de información (palabras clave) |

        ---
        ### Análisis de preguntas abiertas:

        | No | Sub-tareas | Entrada | Salida |
        |----|------------|---------|--------|
        | 6  | Generar análisis de preguntas abiertas | Respuestas abiertas | Clasificación de textos, Tópicos seleccionados |
        | 6.1 | Limpieza de respuestas válidas | Texto de respuestas abiertas | Limpieza de texto, Corrección ortográfica |
        | 6.2 | Clasificar por temas de interés | Respuestas abiertas limpias | Clasificación de textos, Extracción de entidades nombradas |
        | 6.3 | Agrupar comentarios por subtemas | Respuestas organizadas | Agrupamiento de textos (clustering), Tópicos seleccionados |
        | 6.4 | Generar recomendación según sugerencias | Comentarios clasificados | Extracción de información (entidades clave) |

        ---
        | 7  | Generar informe completo de la actividad | Resultados del análisis | Resumen basado en abstracción |
        """, unsafe_allow_html=True)
    
    with st.expander("📊 ¿Qué se busca con el uso de PLN?"):
        st.image("ruta/a/tu/imagen.png", caption="Objetivo del proyecto - PLN", use_column_width=True)

        st.markdown("""
        ## 💡 Objetivo de la solución

        **Comparativa entre el estado actual y el esperado al implementar técnicas de PLN:**

        | Actualmente | Esperado |
        |-------------|-----------|
        | El funcionario dedica 6 horas limpiando y revisando preguntas en formato tabular. | Disminuir significativamente el tiempo de procesamiento gracias a la aplicación de técnicas de PLN. |
        | Procesos **no automatizados** que limitan la eficiencia. | Implementación de **análisis automatizados** y mejora de la productividad. |
        | Exposición a posibles enfermedades laborales debido al alto desgaste visual por el trabajo manual. | Disminución de la exposición a la pantalla, **reducción de riesgo** a enfermedad laboral. |

        ---
        """, unsafe_allow_html=True)
    
# --- Pestaña 2: Entrenamiento ---
with tab2:
    st.header("Entrenar modelo de clasificación de texto")
    archivo_entrenamiento = st.file_uploader("Cargar archivo Excel para entrenamiento", type=["xlsx"], key="entrenamiento")
    if archivo_entrenamiento:
        df_entrenamiento = pd.read_excel(archivo_entrenamiento)
        st.write("Vista previa:", df_entrenamiento.head())

        columnas = df_entrenamiento.columns.tolist()
        variable_texto = st.selectbox("Selecciona la columna de texto", columnas, key="texto_entrenamiento")
        columna_clasificacion = st.selectbox("Selecciona la columna de clasificación", columnas, key="clasificacion_entrenamiento")

        if st.button("Entrenar modelo"):
            accuracy = entrenar_modelo(df_entrenamiento, variable_texto, columna_clasificacion)
            st.success(f"Modelo entrenado con éxito. Accuracy: {accuracy:.4f}")
            st.info("Los archivos 'modelo_entrenado.pkl' y 'label_encoder.pkl' han sido guardados.")


# --- Pestaña 3: Generación de Informe ---
with tab3:
    st.header("Generar informe a partir de datos clasificados")
    archivo_informe = st.file_uploader("Cargar archivo Excel para análisis", type=["xlsx"], key="informe")
    if archivo_informe:
        df_informe = pd.read_excel(archivo_informe)
        st.write("Vista previa:", df_informe.head())

        variable = st.text_input("Nombre de la columna de texto:", key="texto_informe")
        if st.button("Generar Informe"):
            if variable in df_informe.columns:
                df_informe[variable] = df_informe[variable].astype(str)

                predicciones = predecir_clasificacion(df_informe[variable])
                if predicciones is None:
                    st.warning("No se pudo generar el informe debido a la falta del modelo.")
                else:
                    df_informe['clasificacion'] = predicciones

                    resumen = []
                    for clase, grupo in df_informe.groupby('clasificacion'):
                        textos = " ".join(grupo[variable].tolist()).lower()
                        textos = textos.translate(str.maketrans('', '', string.punctuation))
                        palabras = [w for w in textos.split() if w not in stopwords_es and len(w) > 2]
                        mas_comunes = [w for w, _ in Counter(palabras).most_common(10)]
                        resumen.append({
                            'Clase': clase,
                            'Conteo': len(grupo),
                            'Top 10 palabras': ", ".join(mas_comunes)
                        })

                    st.subheader("Resumen por clase")
                    st.dataframe(pd.DataFrame(resumen))

                    recomendaciones_comentarios = df_informe[df_informe['clasificacion'] == 'COMENTARIO'][variable].tolist()
                    recomendaciones_generadas = generar_recomendaciones(recomendaciones_comentarios) if recomendaciones_comentarios else "No se encontraron recomendaciones para analizar."

                    nombre_archivo = "Informe_Encuesta.docx"
                    generar_informe(df_informe[[variable, 'clasificacion']].rename(columns={variable: 'comentario'}), recomendaciones_generadas, nombre_archivo, resumen)

                    with open(nombre_archivo, "rb") as f:
                        st.download_button("📥 Descargar Informe", f, file_name=nombre_archivo)
            else:
                st.error("La columna especificada no existe en el archivo.")

