
import os
import streamlit as st
import pandas as pd
import joblib
from docx import Document
import google.generativeai as genai
from sklearn.base import BaseEstimator, TransformerMixin
import string
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk

# Descargar stopwords si es necesario
nltk.download('stopwords')
from nltk.corpus import stopwords

# Configurar la clave de API de Gemini
genai.configure(api_key=st.secrets("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-1.5-flash")

# Stopwords en español
stopwords_es = set(stopwords.words('spanish'))

# Eliminar clase SpacyVectorizer si ya no se usa
# class SpacyVectorizer(BaseEstimator, TransformerMixin):
#     ...

# Función para predecir nuevas clasificaciones
def predecir_clasificacion(textos):
    pipeline = joblib.load('modelo_rf.pkl')
    le = joblib.load('label_encoder.pkl')
    predicciones = pipeline.predict(textos)
    return le.inverse_transform(predicciones)

# Función para generar recomendaciones con Gemini
def generar_recomendaciones(comentarios):
    entrada = " ".join(comentarios)
    prompt = f"Basado en los siguientes comentarios de recomendación, sugiere mejoras concretas:\n{entrada}\n\nRecomendaciones:"
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Ocurrió un error al generar las recomendaciones: {e}"

# Función para generar informe Word
def generar_informe(comentarios, recomendaciones, nombre_archivo, resumen):
    doc = Document()
    doc.add_heading("Informe de Resultados de la Encuesta", 0)

    doc.add_heading("Resumen por clase", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Clase'
    hdr_cells[1].text = 'Conteo'
    hdr_cells[2].text = 'Top 10 palabras'
    for fila in resumen:
        row_cells = table.add_row().cells
        row_cells[0].text = str(fila['Clase'])
        row_cells[1].text = str(fila['Conteo'])
        row_cells[2].text = fila['Top 10 palabras']

    doc.add_heading("Propuesta de Mejora", level=1)
    doc.add_paragraph(recomendaciones)

    doc.save(nombre_archivo)

# --- Interfaz en Streamlit ---
st.title("Análisis de Encuesta - Ministerio de Defensa")

archivo = st.file_uploader("Cargar archivo Excel", type=["xlsx"])
if archivo:
    df = pd.read_excel(archivo)
    st.write("Vista previa de datos:", df.head())

    variable = st.text_input("Ingrese el nombre de la variable de texto:")
    if st.button("Generar Informe"):
        if variable in df.columns:
            df[variable] = df[variable].astype(str)
            df['clasificacion'] = predecir_clasificacion(df[variable])

            resumen = []
            for clase, grupo in df.groupby('clasificacion'):
                textos = " ".join(grupo[variable].tolist()).lower()
                textos = textos.translate(str.maketrans('', '', string.punctuation))
                palabras = [w for w in textos.split() if w not in stopwords_es and len(w) > 2]
                mas_comunes = [w for w, _ in Counter(palabras).most_common(10)]
                resumen.append({
                    'Clase': clase,
                    'Conteo': len(grupo),
                    'Top 10 palabras': ", ".join(mas_comunes)
                })

            st.subheader("Resumen por clase")
            st.dataframe(pd.DataFrame(resumen))

            recomendaciones_comentarios = df[df['clasificacion'] == 'COMENTARIO'][variable].tolist()
            if recomendaciones_comentarios:
                recomendaciones_generadas = generar_recomendaciones(recomendaciones_comentarios)
            else:
                recomendaciones_generadas = "No se encontraron recomendaciones para analizar."

            nombre_archivo = "Informe_Encuesta.docx"
            generar_informe(df[[variable, 'clasificacion']].rename(columns={variable: 'comentario'}), recomendaciones_generadas, nombre_archivo, resumen)

            with open(nombre_archivo, "rb") as f:
                st.download_button("Descargar Informe", f, file_name=nombre_archivo)
        else:
            st.error("Variable no válida. La columna no existe en el archivo.")
